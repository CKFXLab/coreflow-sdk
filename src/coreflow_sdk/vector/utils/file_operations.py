"""
File Operations Utilities for CoreFlow RAG System

This module provides comprehensive file processing capabilities for RAG systems including:
- File type detection and identification
- Standard file readers for multiple formats (JSON, CSV, PDF, TXT, MD, DOCX, XLSX, HTML)
- Text chunking with configurable parameters
- Metadata extraction and processing
"""

import json
import mimetypes
import hashlib
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass, asdict

# Core dependencies
import pandas as pd

# File processing dependencies (with graceful fallbacks)
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import magic

    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    pass

    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# Import SDK utilities
from ...utils.audit import AppLogger


@dataclass
class FileInfo:
    """Comprehensive file information and metadata."""

    filepath: str
    filename: str
    file_extension: str
    file_type: str
    mime_type: str
    size_bytes: int
    size_mb: float
    created_date: datetime
    modified_date: datetime
    file_hash: str
    encoding: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for metadata storage."""
        return asdict(self)


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""

    text: str
    chunk_id: str
    chunk_index: int
    start_char: int
    end_char: int
    chunk_size: int
    source_file: str
    metadata: Dict[str, Any]

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for storage."""
        return asdict(self)


class FileOperations:
    """Comprehensive file operations for RAG system."""

    def __init__(self):
        """Initialize file operations handler."""
        self.logger = AppLogger(__name__)
        self.supported_formats = {
            "text": [".txt", ".md", ".markdown", ".rst"],
            "json": [".json", ".jsonl"],
            "csv": [".csv", ".tsv"],
            "pdf": [".pdf"],
            "docx": [".docx", ".doc"],
            "xlsx": [".xlsx", ".xls"],
            "html": [".html", ".htm", ".xml"],
        }

        # Initialize mimetypes
        mimetypes.init()

        # Log availability of optional dependencies
        self._log_dependency_status()

    def _log_dependency_status(self):
        """Log the status of optional file processing dependencies."""
        dependencies = {
            "PyPDF2": PDF_AVAILABLE,
            "python-magic": MAGIC_AVAILABLE,
            "python-docx": DOCX_AVAILABLE,
            "beautifulsoup4": HTML_AVAILABLE,
            "openpyxl": XLSX_AVAILABLE,
        }

        available = [name for name, status in dependencies.items() if status]
        unavailable = [name for name, status in dependencies.items() if not status]

        if available:
            self.logger.info(
                f"Available file processing libraries: {', '.join(available)}"
            )
        if unavailable:
            self.logger.warning(
                f"Unavailable file processing libraries: {', '.join(unavailable)}"
            )

    # === FILE TYPE DETECTION ===

    def detect_file_type(self, filepath: str) -> str:
        """
        Detect file type using multiple methods.

        Args:
            filepath: Path to the file

        Returns:
            Detected file type
        """
        filepath = Path(filepath)

        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        # Try extension-based detection first
        extension = filepath.suffix.lower()
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type

        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(str(filepath))
        if mime_type:
            if mime_type.startswith("text/"):
                return "text"
            elif mime_type == "application/json":
                return "json"
            elif mime_type == "text/csv":
                return "csv"
            elif mime_type == "application/pdf":
                return "pdf"
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]:
                return "docx"
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ]:
                return "xlsx"
            elif mime_type in ["text/html", "application/xhtml+xml"]:
                return "html"

        # Try python-magic if available
        if MAGIC_AVAILABLE:
            try:
                file_type = magic.from_file(str(filepath), mime=True)
                if "pdf" in file_type:
                    return "pdf"
                elif "text" in file_type:
                    return "text"
                elif "json" in file_type:
                    return "json"
            except Exception as e:
                self.logger.warning(f"Magic detection failed: {e}")

        # Fallback to text
        self.logger.warning(
            f"Could not determine file type for {filepath}, defaulting to text"
        )
        return "text"

    def get_file_info(self, filepath: str) -> FileInfo:
        """
        Get comprehensive file information and metadata.

        Args:
            filepath: Path to the file

        Returns:
            FileInfo object with complete metadata
        """
        filepath = Path(filepath)

        if not filepath.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        # Get file statistics
        stat = filepath.stat()

        # Calculate file hash
        file_hash = self._calculate_file_hash(filepath)

        # Detect encoding for text files
        encoding = None
        try:
            with open(filepath, "rb") as f:
                raw_data = f.read(1024)
                if raw_data:
                    # Simple encoding detection
                    try:
                        raw_data.decode("utf-8")
                        encoding = "utf-8"
                    except UnicodeDecodeError:
                        try:
                            raw_data.decode("latin-1")
                            encoding = "latin-1"
                        except UnicodeDecodeError:
                            encoding = "unknown"
        except Exception:
            pass

        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(filepath))

        return FileInfo(
            filepath=str(filepath),
            filename=filepath.name,
            file_extension=filepath.suffix.lower(),
            file_type=self.detect_file_type(filepath),
            mime_type=mime_type or "unknown",
            size_bytes=stat.st_size,
            size_mb=round(stat.st_size / (1024 * 1024), 2),
            created_date=datetime.fromtimestamp(stat.st_ctime),
            modified_date=datetime.fromtimestamp(stat.st_mtime),
            file_hash=file_hash,
            encoding=encoding,
        )

    def _calculate_file_hash(self, filepath: Path) -> str:
        """Calculate SHA-256 hash of file."""
        hash_sha256 = hashlib.sha256()
        try:
            with open(filepath, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            self.logger.error(f"Failed to calculate hash for {filepath}: {e}")
            return "unknown"

    # === FILE READERS ===

    def read_text_file(self, filepath: str, encoding: str = "utf-8") -> str:
        """
        Read plain text files (.txt, .md, .rst, etc.).

        Args:
            filepath: Path to the text file
            encoding: Text encoding (default: utf-8)

        Returns:
            File content as string
        """
        try:
            with open(filepath, "r", encoding=encoding) as f:
                content = f.read()
            self.logger.debug(f"Read text file: {filepath} ({len(content)} characters)")
            return content
        except UnicodeDecodeError:
            # Try alternative encodings
            for alt_encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    with open(filepath, "r", encoding=alt_encoding) as f:
                        content = f.read()
                    self.logger.warning(
                        f"Used fallback encoding {alt_encoding} for {filepath}"
                    )
                    return content
                except UnicodeDecodeError:
                    continue
            raise ValueError(
                f"Could not decode file {filepath} with any supported encoding"
            )

    def read_json_file(self, filepath: str) -> Union[Dict, List]:
        """
        Read JSON files (.json, .jsonl).

        Args:
            filepath: Path to the JSON file

        Returns:
            Parsed JSON data
        """
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                if filepath.endswith(".jsonl"):
                    # Handle JSON Lines format
                    data = []
                    for line in f:
                        line = line.strip()
                        if line:
                            data.append(json.loads(line))
                    return data
                else:
                    data = json.load(f)

            self.logger.debug(f"Read JSON file: {filepath}")
            return data
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in file {filepath}: {e}")

    def read_csv_file(self, filepath: str, **kwargs) -> pd.DataFrame:
        """
        Read CSV files (.csv, .tsv).

        Args:
            filepath: Path to the CSV file
            **kwargs: Additional pandas read_csv parameters

        Returns:
            DataFrame with CSV data
        """
        try:
            # Auto-detect separator for TSV files
            if filepath.endswith(".tsv"):
                kwargs.setdefault("sep", "\t")

            df = pd.read_csv(filepath, **kwargs)
            self.logger.debug(
                f"Read CSV file: {filepath} ({len(df)} rows, {len(df.columns)} columns)"
            )
            return df
        except Exception as e:
            raise ValueError(f"Failed to read CSV file {filepath}: {e}")

    def read_pdf_file(self, filepath: str) -> str:
        """
        Read PDF files (.pdf).

        Args:
            filepath: Path to the PDF file

        Returns:
            Extracted text content
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")

        try:
            text_content = []
            with open(filepath, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        self.logger.warning(
                            f"Failed to extract text from page {page_num} in {filepath}: {e}"
                        )

            content = "\n\n".join(text_content)
            self.logger.debug(
                f"Read PDF file: {filepath} ({len(pdf_reader.pages)} pages, {len(content)} characters)"
            )
            return content
        except Exception as e:
            raise ValueError(f"Failed to read PDF file {filepath}: {e}")

    def read_docx_file(self, filepath: str) -> str:
        """
        Read Microsoft Word documents (.docx).

        Args:
            filepath: Path to the DOCX file

        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not available. Install with: pip install python-docx"
            )

        try:
            doc = Document(filepath)
            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            content = "\n\n".join(text_content)
            self.logger.debug(f"Read DOCX file: {filepath} ({len(content)} characters)")
            return content
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file {filepath}: {e}")

    def read_xlsx_file(self, filepath: str, include_sheet_names: bool = True) -> str:
        """
        Read Excel files (.xlsx, .xls).

        Args:
            filepath: Path to the Excel file
            include_sheet_names: Whether to include sheet names in output

        Returns:
            Text representation of Excel data
        """
        if not XLSX_AVAILABLE:
            raise ImportError(
                "openpyxl not available. Install with: pip install openpyxl"
            )

        try:
            # Read all sheets
            excel_data = pd.read_excel(filepath, sheet_name=None)
            text_content = []

            for sheet_name, df in excel_data.items():
                if include_sheet_names:
                    text_content.append(f"Sheet: {sheet_name}")

                # Convert DataFrame to text representation
                if not df.empty:
                    # Include column headers
                    headers = " | ".join(str(col) for col in df.columns)
                    text_content.append(headers)

                    # Include data rows
                    for _, row in df.iterrows():
                        row_text = " | ".join(
                            str(val) for val in row.values if pd.notna(val)
                        )
                        if row_text.strip():
                            text_content.append(row_text)

                text_content.append("")  # Add spacing between sheets

            content = "\n".join(text_content)
            self.logger.debug(
                f"Read XLSX file: {filepath} ({len(excel_data)} sheets, {len(content)} characters)"
            )
            return content
        except Exception as e:
            raise ValueError(f"Failed to read XLSX file {filepath}: {e}")

    def read_html_file(self, filepath: str, extract_text_only: bool = True) -> str:
        """
        Read HTML/XML files (.html, .htm, .xml).

        Args:
            filepath: Path to the HTML file
            extract_text_only: Whether to extract only text content (no tags)

        Returns:
            HTML content or extracted text
        """
        if not HTML_AVAILABLE:
            raise ImportError(
                "beautifulsoup4 not available. Install with: pip install beautifulsoup4"
            )

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            if extract_text_only:
                soup = BeautifulSoup(content, "html.parser")

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Extract text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (
                    phrase.strip() for line in lines for phrase in line.split("  ")
                )
                text = " ".join(chunk for chunk in chunks if chunk)

                self.logger.debug(
                    f"Read HTML file: {filepath} (extracted {len(text)} characters)"
                )
                return text
            else:
                self.logger.debug(
                    f"Read HTML file: {filepath} ({len(content)} characters)"
                )
                return content
        except Exception as e:
            raise ValueError(f"Failed to read HTML file {filepath}: {e}")

    def read_file(self, filepath: str, **kwargs) -> str:
        """
        Universal file reader that automatically detects file type and reads content.

        Args:
            filepath: Path to the file
            **kwargs: Additional parameters for specific readers

        Returns:
            File content as string
        """
        file_type = self.detect_file_type(filepath)

        readers = {
            "text": self.read_text_file,
            "json": lambda path, **kw: json.dumps(self.read_json_file(path), indent=2),
            "csv": lambda path, **kw: self.read_csv_file(path, **kw).to_string(),
            "pdf": self.read_pdf_file,
            "docx": self.read_docx_file,
            "xlsx": self.read_xlsx_file,
            "html": self.read_html_file,
        }

        reader = readers.get(file_type, self.read_text_file)

        try:
            return reader(filepath, **kwargs)
        except Exception as e:
            self.logger.error(f"Failed to read file {filepath} as {file_type}: {e}")
            # Fallback to text reader
            try:
                return self.read_text_file(filepath)
            except Exception as fallback_e:
                raise ValueError(
                    f"Failed to read file {filepath}: {e}. Fallback also failed: {fallback_e}"
                )

    # === TEXT CHUNKING ===

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200,
        source_file: str = "unknown",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[TextChunk]:
        """
        Split text into chunks with configurable size and overlap.

        Args:
            text: Text to chunk
            chunk_size: Maximum characters per chunk
            overlap: Characters to overlap between chunks
            source_file: Source file path for metadata
            metadata: Additional metadata for chunks

        Returns:
            List of TextChunk objects
        """
        if not text or not text.strip():
            return []

        if metadata is None:
            metadata = {}

        chunks = []
        text = text.strip()
        start = 0
        chunk_index = 0

        while start < len(text):
            # Calculate end position
            end = start + chunk_size

            # If this isn't the last chunk, try to find a good breaking point
            if end < len(text):
                # Look for sentence boundaries within the last 100 characters
                search_start = max(start + chunk_size - 100, start)
                sentence_end = -1

                for i in range(end, search_start, -1):
                    if text[i] in ".!?":
                        # Check if this is likely a sentence end (not abbreviation)
                        if i + 1 < len(text) and text[i + 1] in " \n\t":
                            sentence_end = i + 1
                            break

                if sentence_end > 0:
                    end = sentence_end
                else:
                    # Look for paragraph or line breaks
                    for i in range(end, search_start, -1):
                        if text[i] in "\n\r":
                            end = i
                            break
                    else:
                        # Look for word boundaries
                        for i in range(end, search_start, -1):
                            if text[i] in " \t":
                                end = i
                                break

            # Extract chunk text
            chunk_text = text[start:end].strip()

            if chunk_text:
                chunk_id = hashlib.md5(
                    f"{source_file}_{chunk_index}_{start}_{end}".encode()
                ).hexdigest()

                chunk = TextChunk(
                    text=chunk_text,
                    chunk_id=chunk_id,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    chunk_size=len(chunk_text),
                    source_file=source_file,
                    metadata=metadata.copy(),
                )
                chunks.append(chunk)
                chunk_index += 1

            # Move to next chunk with overlap
            if end >= len(text):
                break

            # Calculate next start position with overlap
            start = max(end - overlap, start + 1)

        self.logger.debug(f"Created {len(chunks)} chunks from {len(text)} characters")
        return chunks

    def chunk_by_paragraphs(
        self,
        text: str,
        max_chunk_size: int = 1000,
        source_file: str = "unknown",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[TextChunk]:
        """
        Split text into chunks by paragraphs, respecting max chunk size.

        Args:
            text: Text to chunk
            max_chunk_size: Maximum characters per chunk
            source_file: Source file path for metadata
            metadata: Additional metadata for chunks

        Returns:
            List of TextChunk objects
        """
        if not text or not text.strip():
            return []

        if metadata is None:
            metadata = {}

        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        chunks = []
        current_chunk = ""
        chunk_index = 0
        start_char = 0

        for paragraph in paragraphs:
            # If adding this paragraph would exceed max size, finalize current chunk
            if (
                current_chunk
                and len(current_chunk) + len(paragraph) + 2 > max_chunk_size
            ):
                if current_chunk.strip():
                    end_char = start_char + len(current_chunk)
                    chunk_id = hashlib.md5(
                        f"{source_file}_{chunk_index}_{start_char}_{end_char}".encode()
                    ).hexdigest()

                    chunk = TextChunk(
                        text=current_chunk.strip(),
                        chunk_id=chunk_id,
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=end_char,
                        chunk_size=len(current_chunk.strip()),
                        source_file=source_file,
                        metadata=metadata.copy(),
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                    start_char = end_char

                current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph

        # Add final chunk
        if current_chunk.strip():
            end_char = start_char + len(current_chunk)
            chunk_id = hashlib.md5(
                f"{source_file}_{chunk_index}_{start_char}_{end_char}".encode()
            ).hexdigest()

            chunk = TextChunk(
                text=current_chunk.strip(),
                chunk_id=chunk_id,
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=end_char,
                chunk_size=len(current_chunk.strip()),
                source_file=source_file,
                metadata=metadata.copy(),
            )
            chunks.append(chunk)

        self.logger.debug(
            f"Created {len(chunks)} paragraph-based chunks from {len(paragraphs)} paragraphs"
        )
        return chunks

    # === UTILITY FUNCTIONS ===

    def is_supported_file(self, filepath: str) -> bool:
        """
        Check if file format is supported.

        Args:
            filepath: Path to the file

        Returns:
            True if file format is supported
        """
        extension = Path(filepath).suffix.lower()
        return any(
            extension in extensions for extensions in self.supported_formats.values()
        )

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """
        Get dictionary of supported file formats.

        Returns:
            Dictionary mapping format types to file extensions
        """
        return self.supported_formats.copy()

    def validate_file(self, filepath: str) -> Tuple[bool, str]:
        """
        Validate if file exists and is supported.

        Args:
            filepath: Path to the file

        Returns:
            Tuple of (is_valid, error_message)
        """
        filepath = Path(filepath)

        if not filepath.exists():
            return False, f"File not found: {filepath}"

        if not filepath.is_file():
            return False, f"Path is not a file: {filepath}"

        if not self.is_supported_file(str(filepath)):
            return False, f"Unsupported file format: {filepath.suffix}"

        try:
            # Try to get basic file info
            stat = filepath.stat()
            if stat.st_size == 0:
                return False, f"File is empty: {filepath}"
        except Exception as e:
            return False, f"Cannot access file: {e}"

        return True, "Valid file"

    def extract_metadata_from_content(
        self, content: str, filepath: str = None
    ) -> Dict[str, Any]:
        """
        Extract basic metadata from content.

        Args:
            content: Text content
            filepath: Optional file path for additional context

        Returns:
            Dictionary with extracted metadata
        """
        metadata = {
            "content_length": len(content),
            "word_count": len(content.split()) if content else 0,
            "line_count": content.count("\n") + 1 if content else 0,
            "extraction_date": datetime.now().isoformat(),
        }

        if filepath:
            file_info = self.get_file_info(filepath)
            metadata.update(
                {
                    "source_file": file_info.filename,
                    "file_path": file_info.filepath,
                    "file_type": file_info.file_type,
                    "file_size_mb": file_info.size_mb,
                    "file_modified": file_info.modified_date.isoformat(),
                    "file_hash": file_info.file_hash,
                }
            )

        # Extract first few sentences as potential summary
        if content:
            sentences = content.replace("\n", " ").split(".")[:3]
            summary = ". ".join(
                sentence.strip() for sentence in sentences if sentence.strip()
            )
            if summary and not summary.endswith("."):
                summary += "."
            metadata["auto_summary"] = (
                summary[:200] + "..." if len(summary) > 200 else summary
            )

        return metadata
